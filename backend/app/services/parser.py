import hashlib
import json
from sqlalchemy.orm import Session
from app.models.import_batch import ImportBatch
from app.models.question import Question
from app.models.question_group import QuestionGroup

class ImportError(Exception):
    """Raised when parsing or validation fails during file import."""
    def __init__(self, message: str, report: dict = None):
        super().__init__(message)
        self.report = report or {}

def calculate_file_hash(filepath: str) -> str:
    """Calculate SHA256 hash of a file's binary content."""
    sha256 = hashlib.sha256()
    with open(filepath, "rb") as f:
        while chunk := f.read(8192):
            sha256.update(chunk)
    return sha256.hexdigest()

def calculate_group_hash(g_data: dict) -> str:
    """Calculate SHA256 hash of a QuestionGroup's core content."""
    q_contents = "|".join(f"{q.get('content') or ''}" for q in g_data.get("questions", []))
    content_str = f"{g_data.get('set_id') or ''}|{g_data.get('part')}|{g_data.get('passage_text') or ''}|{g_data.get('audio_url') or ''}|{q_contents}"
    return hashlib.sha256(content_str.encode("utf-8")).hexdigest()

def calculate_question_hash(q_data: dict) -> str:
    """Calculate SHA256 hash of a Question's core content."""
    sorted_options = json.dumps(q_data.get("options") or {}, sort_keys=True)
    content_str = f"{q_data.get('set_id') or ''}|{q_data.get('number') or ''}|{q_data.get('part')}|{q_data.get('type')}|{q_data.get('content') or ''}|{sorted_options}|{q_data.get('reference_answer') or ''}"
    return hashlib.sha256(content_str.encode("utf-8")).hexdigest()

def save_parsed_items(
    db: Session,
    parsed_items: list,
    batch_id: int,
    exam_type: str = "VSTEP_B1",
    language: str = "EN"
) -> dict:
    """Save structured parsed items to database, checking content hash for idempotency."""
    skipped_q_count = 0
    skipped_g_count = 0
    imported_q_count = 0
    imported_g_count = 0

    for item in parsed_items:
        if "questions" in item:
            # It's a Group
            g_hash = calculate_group_hash(item)

            existing_g = db.query(QuestionGroup).filter(
                QuestionGroup.exam_id.is_(None),
                QuestionGroup.content_hash == g_hash
            ).first()

            if existing_g:
                skipped_g_count += 1
                skipped_q_count += len(item["questions"])
                continue

            new_g = QuestionGroup(
                exam_id=None,
                part=item["part"],
                topic=item["topic"],
                passage_text=item["passage_text"],
                audio_url=item["audio_url"],
                image_url=item["image_url"],
                difficulty=item["difficulty"],
                status="draft",
                content_hash=g_hash,
                import_batch_id=batch_id
            )
            db.add(new_g)
            db.commit()
            db.refresh(new_g)
            imported_g_count += 1

            for q_data in item["questions"]:
                q_hash = calculate_question_hash(q_data)

                existing_q = db.query(Question).filter(
                    Question.exam_id.is_(None),
                    Question.content_hash == q_hash
                ).first()

                if existing_q:
                    skipped_q_count += 1
                    continue

                new_q = Question(
                    exam_id=None,
                    group_id=new_g.id,
                    part=q_data["part"],
                    type=q_data["type"],
                    content=q_data["content"],
                    audio_url=q_data["audio_url"],
                    image_url=q_data["image_url"],
                    options=q_data["options"],
                    reference_answer=q_data["reference_answer"],
                    difficulty=q_data["difficulty"],
                    clo=q_data["clo"],
                    topic=q_data["topic"],
                    explanation=q_data["explanation"],
                    status="draft",
                    content_hash=q_hash,
                    import_batch_id=batch_id,
                    exam_type=exam_type,
                    language=language
                )
                db.add(new_q)
                imported_q_count += 1
            db.commit()

        else:
            # It's a Standalone Question
            q_hash = calculate_question_hash(item)

            existing_q = db.query(Question).filter(
                Question.exam_id.is_(None),
                Question.content_hash == q_hash
            ).first()

            if existing_q:
                skipped_q_count += 1
                continue

            new_q = Question(
                exam_id=None,
                group_id=None,
                part=item["part"],
                type=item["type"],
                content=item["content"],
                audio_url=item["audio_url"],
                image_url=item["image_url"],
                options=item["options"],
                reference_answer=item["reference_answer"],
                difficulty=item["difficulty"],
                clo=item["clo"],
                topic=item["topic"],
                explanation=item["explanation"],
                status="draft",
                content_hash=q_hash,
                import_batch_id=batch_id,
                exam_type=exam_type,
                language=language
            )
            db.add(new_q)
            db.commit()
            imported_q_count += 1

    return {
        "imported_questions": imported_q_count,
        "imported_groups": imported_g_count,
        "skipped_questions": skipped_q_count,
        "skipped_groups": skipped_g_count
    }

def convert_doc_to_docx(filepath: str, out_dir: str = None) -> str:
    """
    Convert a legacy Word document (.doc) to XML format (.docx) using headless LibreOffice.
    
    - If the input file already has a .docx extension, returns the input filepath unchanged (passthrough).
    - If the input file has a .doc extension, converts it using `soffice --headless --convert-to docx`.
    - Caches result: if the target .docx file already exists, skips conversion and returns its path immediately.
    - If soffice is missing from PATH, raises a clear, helpful FileNotFoundError.
    """
    import os
    import shutil
    import subprocess
    
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Source file not found: {filepath}")
        
    # Normalize path extension
    _, ext = os.path.splitext(filepath.lower())
    
    # 1. Passthrough: if already docx, return path directly
    if ext == ".docx":
        return filepath
        
    if ext != ".doc":
        raise ValueError(f"Unsupported file format for conversion: {ext} (only .doc is supported)")
        
    # Determine output directory (default to input file's directory)
    if out_dir is None:
        out_dir = os.path.dirname(os.path.abspath(filepath))
    else:
        out_dir = os.path.abspath(out_dir)
        os.makedirs(out_dir, exist_ok=True)
        
    # Determine target docx filename
    base = os.path.basename(filepath)
    stem, _ = os.path.splitext(base)
    target_docx = os.path.join(out_dir, f"{stem}.docx")
    
    # 2. Cache check: if target docx exists, skip conversion
    if os.path.exists(target_docx):
        return target_docx
        
    # 3. Tool presence check
    soffice_path = shutil.which("soffice")
    if soffice_path is None:
        # Check common windows installation folder fallback just in case
        common_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
        ]
        for cp in common_paths:
            if os.path.exists(cp):
                soffice_path = cp
                break
                
    if soffice_path is None:
        raise FileNotFoundError(
            "LibreOffice 'soffice' executable not found on PATH or standard directories.\n"
            "Please install LibreOffice (https://www.libreoffice.org) and make sure 'soffice' "
            "is added to your system environment variables (PATH)."
        )
        
    # 4. Conversion via subprocess (using list of args, no shell=True)
    cmd = [
        soffice_path,
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        out_dir,
        filepath
    ]
    
    try:
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=True)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"LibreOffice conversion failed for {filepath}.\n"
            f"Exit code: {e.returncode}\n"
            f"Error output: {e.stderr or e.stdout}"
        ) from e
        
    if not os.path.exists(target_docx):
        raise RuntimeError(
            f"LibreOffice command finished successfully, but target file was not created: {target_docx}"
        )
        
    return target_docx


def parse_b1_reading_docx(filepath: str) -> dict:
    """
    Parses a B1 Reading and Writing test .docx file.
    Returns:
        dict: {
            "set_id": "EB12601",
            "items": [
                {
                    "number": int,
                    "part": int (1-6),
                    "section": int,
                    "type": str ('choice', 'fill', 'writing'),
                    "content": str,
                    "options": dict
                },
                ...
            ]
        }
    """
    import os
    import re
    import docx
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    doc = docx.Document(filepath)

    # 1. Extract set_id
    set_id = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                m = re.search(r'(?i)EB1\.?\s*(\d+)', cell.text)
                if m:
                    set_id = f"EB1{m.group(1)}"
                    break
            if set_id:
                break
        if set_id:
            break

    if not set_id:
        for p in doc.paragraphs:
            m = re.search(r'(?i)EB1\.?\s*(\d+)', p.text)
            if m:
                set_id = f"EB1{m.group(1)}"
                break

    if not set_id:
        filename = os.path.basename(filepath)
        m = re.search(r'(?i)EB1\.?\s*(\d+)', filename)
        if m:
            set_id = f"EB1{m.group(1)}"
        else:
            set_id = "EB19999"

    # 2. Traverse body elements
    body_elements = list(doc.element.body)

    current_part = None  # "R" or "W"
    current_section = None  # 1-4 for Reading, 1-2 for Writing
    items = []

    s3_passage_paragraphs = []
    s4_passage_paragraphs = []
    w1_paragraphs = []
    w2_paragraphs = []

    visited_indices = set()

    def extract_inline_options_4(text):
        m = re.match(r'^\s*A[\.\)]\s*(.*?)\s+B[\.\)]\s*(.*?)\s+C[\.\)]\s*(.*?)\s+D[\.\)]\s*(.*)$', text)
        if m:
            return {
                "A": m.group(1).strip(),
                "B": m.group(2).strip(),
                "C": m.group(3).strip(),
                "D": m.group(4).strip()
            }
        pos_a = text.find("A.")
        pos_b = text.find("B.")
        pos_c = text.find("C.")
        pos_d = text.find("D.")
        if pos_a != -1 and pos_b != -1 and pos_c != -1 and pos_d != -1:
            return {
                "A": text[pos_a+2:pos_b].strip(),
                "B": text[pos_b+2:pos_c].strip(),
                "C": text[pos_c+2:pos_d].strip(),
                "D": text[pos_d+2:].strip()
            }
        return {}

    i = 0
    while i < len(body_elements):
        if i in visited_indices:
            i += 1
            continue
            
        child = body_elements[i]
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            text = p.text.strip()
            
            # Check part/section transitions
            if re.search(r'(?i)PART\s*ONE|READING', text):
                current_part = "R"
                i += 1
                continue
            elif re.search(r'(?i)PART\s*TWO|WRITING', text):
                current_part = "W"
                i += 1
                continue
                
            sec_m = re.search(r'(?i)Section\s+(\d+)', text)
            if sec_m:
                sec_num = int(sec_m.group(1))
                if current_part == "R":
                    current_section = sec_num
                elif current_part == "W":
                    current_section = sec_num
                i += 1
                continue
                
            # Parse based on current section
            if current_part == "R":
                if current_section == 1:
                    # S1 Q1-10
                    q_m = re.match(r'^\s*(\d+)\.\s*(.*)', text, re.DOTALL)
                    if q_m:
                        q_num = int(q_m.group(1))
                        q_content = q_m.group(2).strip()
                        options = {}
                        j = i + 1
                        while j < len(body_elements):
                            next_child = body_elements[j]
                            if isinstance(next_child, CT_P):
                                next_p = Paragraph(next_child, doc)
                                next_text = next_p.text.strip()
                                if next_text:
                                    options = extract_inline_options_4(next_text)
                                    visited_indices.add(j)
                                    break
                            j += 1
                        items.append({
                            "number": q_num,
                            "part": 1,
                            "section": 1,
                            "type": "choice",
                            "content": q_content,
                            "options": options
                        })
                        
                elif current_section == 3:
                    # S3 Q16-20
                    q_m = re.match(r'^\s*(\d+)\.\s*(.*)', text, re.DOTALL)
                    if q_m:
                        q_num = int(q_m.group(1))
                        q_content = q_m.group(2).strip()
                        options = {}
                        j = i + 1
                        opt_count = 0
                        while j < len(body_elements) and opt_count < 4:
                            next_child = body_elements[j]
                            if isinstance(next_child, CT_P):
                                next_p = Paragraph(next_child, doc)
                                next_text = next_p.text.strip()
                                if next_text:
                                    opt_m = re.match(r'^\s*([A-D])[\.\)]\s*(.*)', next_text)
                                    if opt_m:
                                        options[opt_m.group(1)] = opt_m.group(2).strip()
                                        visited_indices.add(j)
                                        opt_count += 1
                                    else:
                                        break
                                else:
                                    j += 1
                                    continue
                            else:
                                break
                            j += 1
                            
                        items.append({
                            "number": q_num,
                            "part": 3,
                            "section": 3,
                            "type": "choice",
                            "content": q_content,
                            "options": options
                        })
                    else:
                        if text and not text.startswith("Read the text"):
                            s3_passage_paragraphs.append(text)
                            
                elif current_section == 4:
                    # S4 Q21-30
                    if text and not text.startswith("Read the text") and not text.startswith("THE GORILLA"):
                        s4_passage_paragraphs.append(text)
                        
            elif current_part == "W":
                if current_section == 1:
                    # W1 (rewriting)
                    if text and not text.startswith("Finish each of the following sentences") and not text.startswith("Example") and not text.startswith("Answer"):
                        w1_paragraphs.append(text)
                elif current_section == 2:
                    # W2 (essay/letter)
                    if text and not text.startswith("Now write a letter"):
                        w2_paragraphs.append(text)
                        
        elif isinstance(child, CT_Tbl):
            t = Table(child, doc)
            if current_part == "R" and current_section == 2:
                # S2 Q11-15 table
                for row in t.rows:
                    cells = row.cells
                    if len(cells) >= 3:
                        cell0_text = cells[0].text.strip()
                        num_match = re.match(r'^\s*(\d+)\.?\s*$', cell0_text)
                        if num_match:
                            q_num = int(num_match.group(1))
                            if 11 <= q_num <= 15:
                                content = cells[1].text.strip()
                                options_text = cells[2].text.strip()
                                options = {}
                                pos_a = options_text.find("A.")
                                pos_b = options_text.find("B.")
                                pos_c = options_text.find("C.")
                                if pos_a != -1 and pos_b != -1 and pos_c != -1:
                                    options = {
                                        "A": options_text[pos_a+2:pos_b].strip(),
                                        "B": options_text[pos_b+2:pos_c].strip(),
                                        "C": options_text[pos_c+2:].strip()
                                    }
                                else:
                                    for line in options_text.split("\n"):
                                        line = line.strip()
                                        if line:
                                            m_opt = re.match(r'^\s*([A-C])\.\s*(.*)', line)
                                            if m_opt:
                                                options[m_opt.group(1)] = m_opt.group(2).strip()
                                items.append({
                                    "number": q_num,
                                    "part": 2,
                                    "section": 2,
                                    "type": "choice",
                                    "content": content,
                                    "options": options
                                })
                                
        i += 1

    # Post-process S4 (fill blanks)
    s4_text = "\n".join(s4_passage_paragraphs).replace("\n", " ")
    s4_text = re.sub(r'\s+', ' ', s4_text)
    sentences = [s.strip() for s in re.split(r'(?<!\.)\.(?!\.)(?=\s+[A-Z]|\s*\(\d+\)|\s*$)', s4_text) if s.strip()]

    for q_num in range(21, 31):
        blank_str = f"({q_num})"
        content = ""
        for s in sentences:
            if blank_str in s:
                content = s + "."
                break
        if not content:
            content = s4_text
            
        items.append({
            "number": q_num,
            "part": 4,
            "section": 4,
            "type": "fill",
            "content": content,
            "options": {}
        })

    # Add W1 and W2 items
    items.append({
        "number": 31,
        "part": 5,
        "section": 1,
        "type": "writing",
        "content": "\n".join(w1_paragraphs),
        "options": {}
    })

    items.append({
        "number": 32,
        "part": 6,
        "section": 2,
        "type": "writing",
        "content": "\n".join(w2_paragraphs),
        "options": {}
    })

    # Sort items by question number
    items.sort(key=lambda x: x["number"])

    return {
        "set_id": set_id,
        "items": items
    }


def parse_b1_answer_key(filepath: str) -> dict:
    """
    Parses a B1 Reading and Writing answer key .docx file.
    Returns:
        dict: {
            1: "A",
            ...
            21: "fact",
            ...
        }
    """
    import os
    import re
    import docx
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    doc = docx.Document(filepath)

    in_writing = False
    answers = {}

    for child in doc.element.body:
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            text = p.text.strip()
            if re.search(r'(?i)PART\s*2|WRITING', text):
                in_writing = True
        elif isinstance(child, CT_Tbl) and not in_writing:
            t = Table(child, doc)
            
            question_cols = []
            header_row_idx = None
            
            for r_idx, row in enumerate(t.rows):
                row_texts = [c.text.strip() for c in row.cells]
                if "Câu" in row_texts and "Đáp án" in row_texts:
                    header_row_idx = r_idx
                    for c_idx, val in enumerate(row_texts):
                        if val == "Câu":
                            question_cols.append(c_idx)
                    break
                    
            if header_row_idx is not None:
                for r_idx in range(header_row_idx + 1, len(t.rows)):
                    row = t.rows[r_idx]
                    for c_idx in question_cols:
                        if c_idx + 1 < len(row.cells):
                            q_num_text = row.cells[c_idx].text.strip()
                            q_ans_text = row.cells[c_idx + 1].text.strip()
                            if re.match(r'^\d+$', q_num_text):
                                num = int(q_num_text)
                                if 1 <= num <= 30:
                                    if 1 <= num <= 20:
                                        answers[num] = q_ans_text.strip().upper()
                                    else:
                                        answers[num] = q_ans_text.strip()
                                        
    return answers


def import_b1_reading_set(db: Session, docx_path: str, key_path: str) -> dict:
    """
    Parses, merges answers, validates, and imports a real-format B1 Reading and Writing set
    into the Question Bank with status='draft', exam_type='VSTEP_B1', language='EN'.
    """
    # 1. Parse docx
    docx_res = parse_b1_reading_docx(docx_path)
    set_id_docx = docx_res["set_id"]
    items = docx_res["items"]

    # 2. Parse answer key
    answers = parse_b1_answer_key(key_path)

    # 3. Merge and validate answers
    errors = []
    for item in items:
        item["set_id"] = set_id_docx
        item.setdefault("difficulty", "medium")
        item.setdefault("explanation", None)
        item.setdefault("clo", None)
        item.setdefault("topic", None)
        item.setdefault("audio_url", None)
        item.setdefault("image_url", None)

        q_num = item["number"]
        location_q = f"Question {q_num}"
        if item["type"] != "writing":
            if q_num not in answers:
                errors.append({
                    "location": location_q,
                    "type": "missing_answer",
                    "message": f"No answer found in KEY for Question {q_num}."
                })
            else:
                ans = answers[q_num]
                item["reference_answer"] = ans
                if item["type"] == "choice" and item["options"]:
                    if ans not in item["options"]:
                        errors.append({
                            "location": location_q,
                            "type": "invalid_answer",
                            "message": f"Invalid answer '{ans}' in KEY for Question {q_num}."
                        })
        else:
            item["reference_answer"] = None

    if errors:
        raise ImportError(f"Validation failed during merging for {docx_path}", {
            "file": docx_path,
            "errors": errors
        })

    # 4. Create ImportBatch (Atomic database transactions)
    file_hash = calculate_file_hash(docx_path)
    
    batch = ImportBatch(
        source_file=docx_path,
        content_hash=file_hash,
        status="imported"
    )
    db.add(batch)
    db.commit()
    db.refresh(batch)

    try:
        result = save_parsed_items(db, items, batch.id, exam_type="VSTEP_B1", language="EN")
        return result
    except Exception as e:
        db.rollback()
        raise e


def parse_b1_listening_docx(filepath: str) -> dict:
    """
    Parses a B1 Listening test .docx file.
    Returns:
        dict: {
            "set_id": "LB12601",
            "items": [...]
        }
    """
    import os
    import re
    import docx

    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    doc = docx.Document(filepath)

    # 1. Extract set_id
    set_id = None

    def find_set_id(text):
        if not text:
            return None
        text_clean = text.strip()
        m = re.search(r'(?i)LB1[-.\s]*(\d+)', text_clean)
        if m:
            return f"LB1{m.group(1)}"
        m = re.search(r'(?i)Mã\s*đề(?:\s*thi)?\s*[:.-]?\s*(?:LB1[-.\s]*)?(\d+)', text_clean)
        if m:
            return f"LB1{m.group(1)}"
        return None
    
    # Try header/footer first
    for section in doc.sections:
        for f in (section.footer, section.first_page_footer, section.even_page_footer,
                  section.header, section.first_page_header, section.even_page_header):
            if f and f.paragraphs:
                for p in f.paragraphs:
                    set_id = find_set_id(p.text)
                    if set_id:
                        break
            if set_id:
                break
        if set_id:
            break

    # Try paragraphs
    if not set_id:
        for p in doc.paragraphs:
            set_id = find_set_id(p.text)
            if set_id:
                break

    # Try tables
    if not set_id:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    set_id = find_set_id(cell.text)
                    if set_id:
                        break
                if set_id:
                    break
            if set_id:
                break

    # Try filename fallback
    if not set_id:
        filename = os.path.basename(filepath)
        set_id = find_set_id(filename)
        if not set_id:
            m = re.search(r'\d{4}', filename)
            if m:
                set_id = f"LB1{m.group(0)}"
            else:
                set_id = "LB19999"

    items = []

    # Helper to find all image shape indices in doc.inline_shapes for a cell
    def get_cell_image_indices(cell):
        indices = []
        drawings = cell._tc.xpath('.//w:drawing')
        for d in drawings:
            blips = d.xpath('.//a:blip/@r:embed') or d.xpath('.//a:blip/@r:link')
            if blips:
                rId = blips[0]
                for idx, shape in enumerate(doc.inline_shapes):
                    shape_rId = shape._inline.xpath('.//a:blip/@r:embed') or shape._inline.xpath('.//a:blip/@r:link')
                    if shape_rId and shape_rId[0] == rId:
                        indices.append(idx)
                        break
        return indices

    # Table 0: Q1-5 (Choice questions)
    image_report = []
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        for i in range(5):
            if 2 * i + 1 < len(t0.rows):
                text_cell = t0.rows[2*i].cells[0]
                img_cell = t0.rows[2*i+1].cells[0]

                text = text_cell.text.strip()
                content = re.sub(r'^\d+\.\s*', '', text)

                img_indices = get_cell_image_indices(img_cell)
                image_report.append(f"Q{i+1}: {len(img_indices)} images")

                img_url = ",".join(map(str, img_indices)) if img_indices else None

                items.append({
                    "number": i + 1,
                    "part": 1,
                    "section": 1,
                    "type": "choice",
                    "content": content,
                    "options": {"A": "", "B": "", "C": ""},
                    "reference_answer": None,
                    "image_url": img_url,
                    "audio_url": None,
                    "difficulty": "medium",
                    "clo": None,
                    "topic": None,
                    "explanation": None
                })
    if image_report:
        print(f"[REPORT] B1 Listening Section 1 image count report: {'; '.join(image_report)}")

    # Table 1: Q6-11 (Fill questions)
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        cell_text = t1.rows[0].cells[0].text
        for line in cell_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            m = re.search(r'\((\d+)\)', line)
            if m:
                q_num = int(m.group(1))
                items.append({
                    "number": q_num,
                    "part": 2,
                    "section": 2,
                    "type": "fill",
                    "content": line,
                    "options": {},
                    "reference_answer": None,
                    "image_url": None,
                    "audio_url": None,
                    "difficulty": "medium",
                    "clo": None,
                    "topic": None,
                    "explanation": None
                })

    # Table 2: Q12-15 (Fill questions)
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        cell_text = t2.rows[0].cells[0].text
        for line in cell_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            m = re.search(r'\((\d+)\)', line)
            if m:
                q_num = int(m.group(1))
                items.append({
                    "number": q_num,
                    "part": 2,
                    "section": 2,
                    "type": "fill",
                    "content": line,
                    "options": {},
                    "reference_answer": None,
                    "image_url": None,
                    "audio_url": None,
                    "difficulty": "medium",
                    "clo": None,
                    "topic": None,
                    "explanation": None
                })

    items.sort(key=lambda x: x["number"])
    return {
        "set_id": set_id,
        "items": items
    }


def parse_b1_listening_key(filepath: str) -> dict:
    """
    Parses a B1 Listening Answer Key file (doc/docx).
    Returns:
        dict: {number: answer}
    """
    import os
    import re
    import docx

    docx_path = convert_doc_to_docx(filepath)
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Converted file not found: {docx_path}")

    doc = docx.Document(docx_path)
    answers = {}

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for line in cell.text.split('\n'):
                    line = line.strip()
                    m = re.match(r'^(\d+)\.\s*(.*)$', line)
                    if m:
                        num = int(m.group(1))
                        ans = m.group(2).strip()
                        if 1 <= num <= 5:
                            answers[num] = ans.upper()
                        else:
                            answers[num] = ans
    return answers


def parse_b1_speaking_card(filepath: str, set_id: str) -> list:
    """
    Parses a B1 Speaking Card file (doc/docx) and extracts the 3 parts for the given set_id.
    Returns:
        list of dicts
    """
    import os
    import re
    import docx

    clean_id = re.sub(r'(?i)^[LE]B1[-.]?', '', set_id)
    m_target = re.search(r'\d+', clean_id)
    if not m_target:
        m_target = re.search(r'\d+', set_id)
    if not m_target:
        raise ValueError(f"Could not extract digits from set_id: {set_id}")
    target_digits = m_target.group(0)

    docx_path = convert_doc_to_docx(filepath)
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Converted file not found: {docx_path}")

    doc = docx.Document(docx_path)

    card_paragraphs = []
    found_card = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        m_card = re.search(r'(?i)Speaking\s*card\s*B1\s*(\d+)', text)
        if m_card:
            card_digits = m_card.group(1)
            if card_digits == target_digits:
                found_card = True
                continue
            else:
                found_card = False

        if found_card:
            if text.startswith('---') or re.search(r'(?i)Speaking\s*card\s*B1', text):
                found_card = False
            else:
                card_paragraphs.append(text)

    # Search within tables if not found in paragraphs
    if not card_paragraphs:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    lines = [line.strip() for line in cell.text.split('\n') if line.strip()]
                    for idx, line in enumerate(lines):
                        m_card = re.search(r'(?i)Speaking\s*card\s*B1\s*(\d+)', line)
                        if m_card and m_card.group(1) == target_digits:
                            for k in range(idx + 1, len(lines)):
                                next_line = lines[k]
                                if next_line.startswith('---') or re.search(r'(?i)Speaking\s*card\s*B1', next_line):
                                    break
                                card_paragraphs.append(next_line)
                            break
                    if card_paragraphs:
                        break
                if card_paragraphs:
                    break
            if card_paragraphs:
                break

    part1_content = ""
    part2_content = ""
    part3_content = ""

    for p_text in card_paragraphs:
        if re.match(r'(?i)^Part\s*1', p_text):
            part1_content = p_text
        elif re.match(r'(?i)^Part\s*2', p_text):
            part2_content = p_text
        elif re.match(r'(?i)^Part\s*3', p_text):
            part3_content = p_text

    if not part1_content and len(card_paragraphs) >= 3:
        part1_content = card_paragraphs[0]
        part2_content = card_paragraphs[1]
        part3_content = card_paragraphs[2]

    prompts = [part1_content, part2_content, part3_content]
    items = []
    for idx, prompt in enumerate(prompts):
        items.append({
            "number": idx + 1,
            "part": idx + 1,
            "section": idx + 1,
            "type": "speaking",
            "content": prompt or f"Part {idx + 1} prompt",
            "options": {},
            "reference_answer": None,
            "image_url": None,
            "audio_url": None,
            "difficulty": "medium",
            "clo": None,
            "topic": None,
            "explanation": None
        })
    return items

