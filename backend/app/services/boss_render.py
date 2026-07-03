"""Helper render-từ-JSON THAM CHIẾU (SPEC-FACTORY-010) — de-risk D2 Option (a).

Nhà máy `boss_factory` xuất item JSON đúng shape ngân hàng của đối tác. Module NÀY chứng minh
item đó ĐỦ thông tin để render thành đề đọc-được (GV ký được), và là SPEC MẪU để đối tác copy
nhánh "render item từ JSON" vào pipeline của họ — KHÔNG thay pipeline sếp, chỉ tham chiếu.

Nhận 1 bundle (từ `boss_factory.export_*_bundle`, mọi skill) → 1 file .docx gồm PHẦN ĐỀ +
PHẦN ĐÁP ÁN (GV soát/ký). Dùng python-docx (đã có ở requirements base — 0 dep mới). Part ẢNH
(R2 dạng ảnh / L1 chọn-tranh) chèn placeholder '[ẢNH ...]' tham chiếu (asset giao riêng).

PHÒNG THỦ (review S54): renderer nhận cả bundle JSON méo (đọc từ đĩa qua render_bundle_file) →
mọi vòng lặp có guard isinstance (points/options/questions/s4_raw/l1 non-list-or-dict → bỏ qua,
KHÔNG crash cả file).
"""
import json
import logging
import os

logger = logging.getLogger(__name__)

R1_OPTS = ["A", "B", "C", "D"]
R2_OPTS = ["A", "B", "C"]


def _as_list(x) -> list:
    return x if isinstance(x, list) else []


def _flat(x) -> str:
    if isinstance(x, list):
        return "; ".join(str(i) for i in x)
    return str(x or "")


def _hdr(doc, item: dict, idx: int, label: str):
    """Tiêu đề 1 câu/nhóm (in đậm) kèm độ khó / nguồn seed / cờ QC cho GV."""
    meta = []
    if item.get("do_kho"):
        meta.append(f"độ khó {item['do_kho']}")
    if item.get("nguon_seed"):
        meta.append(f"nguồn {item['nguon_seed']}")
    if item.get("qc_ok") is False:
        meta.append("⚠ QC: " + "; ".join(item.get("qc_issues") or []))
    p = doc.add_paragraph()
    p.add_run(f"{label} {idx}" + (f"  [{' · '.join(meta)}]" if meta else "")).bold = True


def _opts(doc, options, keys):
    if not isinstance(options, dict):        # phòng thủ: options méo (list/None) → bỏ qua, không crash
        return
    for k in keys:
        if k in options:
            doc.add_paragraph(f"    {k}. {options[k]}")


def _render_r1(doc, idx, it, ak):
    q = it.get("s1_item") or {}
    _hdr(doc, it, idx, "Câu")
    doc.add_paragraph(q.get("stem", ""))
    _opts(doc, q.get("options"), R1_OPTS)
    ak.append(f"Câu {idx}: {q.get('answer', '')}")


def _render_r2(doc, idx, it, ak):
    q = it.get("s2_item") or {}
    _hdr(doc, it, idx, "Câu (thông báo)")
    doc.add_paragraph(q.get("stem", ""))
    _opts(doc, q.get("options"), R2_OPTS)
    ak.append(f"Câu {idx}: {q.get('answer', '')}")


def _render_r3(doc, idx, it, ak):
    s3 = it.get("s3_item") or {}
    questions = _as_list(s3.get("questions"))
    ans_map = it.get("s3_answers") if isinstance(it.get("s3_answers"), dict) else {}
    # số câu THEO ĐỀ THẬT (VSTEP R3 = 16-20), lấy từ s3_answers nếu có; else 16+j → key KHỚP đề GV thấy
    nums = (sorted(ans_map.keys(), key=lambda x: int(x)) if ans_map
            else [str(16 + j) for j in range(len(questions))])
    _hdr(doc, it, idx, "Nhóm đọc-hiểu")
    doc.add_paragraph(s3.get("passage", ""))
    marks = []
    for j, q in enumerate(questions):
        q = q if isinstance(q, dict) else {}
        num = nums[j] if j < len(nums) else str(16 + j)
        doc.add_paragraph(f"  {num}. {q.get('stem', '')}")
        _opts(doc, q.get("options"), R1_OPTS)
        marks.append(f"{num}-{ans_map.get(num) or q.get('answer', '')}")
    ak.append(f"Nhóm {idx}: " + ", ".join(marks))


def _render_r4(doc, idx, it, ak):
    s4 = it.get("s4_item") or {}
    _hdr(doc, it, idx, "Bài điền-từ (cloze)")
    for blk in _as_list(s4.get("s4_raw")):
        blk = blk if isinstance(blk, dict) else {}
        if blk.get("kind") != "tbl":              # hộp từ render riêng từ word_box cho gọn
            doc.add_paragraph(blk.get("text", ""))
    box = it.get("word_box")
    if not (isinstance(box, list) and box):       # word_box rỗng → suy từ đáp án để GV vẫn thấy hộp từ
        box = sorted({str(v) for v in (s4.get("s4_answers") or {}).values()})
    if box:
        doc.add_paragraph("Hộp từ: " + ", ".join(str(w) for w in box))
    ans = it.get("answers") or s4.get("s4_answers") or {}
    ak.append(f"Cloze {idx}: " + ", ".join(f"{k}={v}" for k, v in ans.items()))


def _render_w1(doc, idx, it, ak):
    w = it.get("w1_item") or {}
    _hdr(doc, it, idx, "Câu viết-lại (giữ nghĩa)")
    doc.add_paragraph(w.get("original", ""))
    doc.add_paragraph("→ " + (w.get("prompt") or ""))
    ak.append(f"Câu {idx}: {w.get('answer', '')}")


def _render_w2(doc, idx, it, ak):
    w = it.get("w2_item") or {}
    _hdr(doc, it, idx, "Đề viết thư (~100 từ)")
    for key in ("role", "situation"):
        if w.get(key):
            doc.add_paragraph(w[key])
    for pt in _as_list(w.get("points")):          # guard: points méo (scalar) → bỏ qua, không crash
        doc.add_paragraph(f"  • {pt}")
    if w.get("instruction"):
        doc.add_paragraph(w["instruction"])
    ak.append(f"Đề {idx}: (tự luận — chấm theo tiêu chí, không đáp án cố định)")


def _render_speak_part(doc, label, val):
    """Part 1/Part 3: list câu hỏi → từng bullet (nhất quán w1 points/s3 questions); scalar → 1 dòng."""
    if isinstance(val, list):
        doc.add_paragraph(f"{label}:")
        for q in val:
            doc.add_paragraph(f"  • {q}")
    else:
        doc.add_paragraph(f"{label}: " + str(val or ""))


def _render_speak(doc, idx, it, ak):
    c = it.get("speak_card") or {}
    _hdr(doc, it, idx, f"Thẻ Nói ({c.get('code', '')})")
    _render_speak_part(doc, "Part 1", c.get("part1"))
    doc.add_paragraph("Part 2 (chủ đề): " + _flat(c.get("part2_topic")))
    _render_speak_part(doc, "Part 3", c.get("part3"))
    ak.append(f"Thẻ {idx}: (Nói — chấm theo tiêu chí, không đáp án)")


def _render_lis(doc, idx, it, ak):
    li = it.get("lis_item") or {}
    tr = it.get("transcripts") or {}
    _hdr(doc, it, idx, f"Bài Nghe ({li.get('code', '')})")
    doc.add_paragraph("PART 1 — Chọn tranh (5 câu hội thoại, đọc 2 lần):")
    for j, q in enumerate(_as_list(tr.get("l1")), 1):
        q = q if isinstance(q, dict) else {}
        imgs = str(q.get("image_urls") or "").strip()      # tên file ảnh thật (FACTORY-011) nếu đã sinh
        has_img = bool(imgs) and any(part.strip() for part in imgs.split(","))   # guard: bỏ chuỗi ',,' lệch
        img_ref = imgs.replace(",", " / ") if has_img else "A/B/C — asset giao riêng"
        doc.add_paragraph(f"  Câu {j}. {q.get('stem', '')}   [ẢNH: {img_ref}]")
        _opts(doc, q.get("options"), R2_OPTS)
        doc.add_paragraph(f"     (Kịch bản: {q.get('transcript', '')})")
    n_l2 = li.get("l2_count") if isinstance(li.get("l2_count"), int) else 10
    doc.add_paragraph(f"PART 2 — Điền từ (nghe monologue, câu 6-{5 + n_l2}):")
    doc.add_paragraph("  (Kịch bản: " + str(tr.get("l2") or "") + ")")
    ans = li.get("answers") or {}
    ak.append(f"Nghe {idx}: " + ", ".join(f"{k}={v}" for k, v in ans.items()))


_RENDERERS = {
    "reading_s1": _render_r1,
    "reading_s2_notice": _render_r2,
    "reading_s3_comprehension": _render_r3,
    "reading_s4_cloze": _render_r4,
    "writing_w1_rewrite": _render_w1,
    "writing_w2_letter": _render_w2,
    "speaking": _render_speak,
    "listening": _render_lis,
}

_SKILL_LABEL = {
    "reading_s1": "Đọc R1", "reading_s2_notice": "Đọc R2", "reading_s3_comprehension": "Đọc R3",
    "reading_s4_cloze": "Đọc R4 (cloze)", "writing_w1_rewrite": "Viết W1", "writing_w2_letter": "Viết W2",
    "speaking": "Nói", "listening": "Nghe",
}


def render_bundle_docx(bundle: dict, out_path: str) -> str:
    """Render 1 bundle (export_*_bundle) → .docx THAM CHIẾU: phần đề + phần đáp án (GV soát/ký).

    Dispatcher theo bundle['skill']; skill lạ → dump field cơ bản (không mất dữ liệu). Trả đường dẫn."""
    from docx import Document  # noqa: PLC0415 — python-docx là base dep; lazy để import module không phụ thuộc SDK
    skill = str(bundle.get("skill") or "?")
    renderer = _RENDERERS.get(skill)
    doc = Document()
    doc.add_heading(f"Bàn giao {_SKILL_LABEL.get(skill, skill)} — {bundle.get('spec', '')}", 0)
    note_text = str(bundle.get("note") or "").strip()          # strip: note toàn khoảng trắng → bỏ
    if note_text:
        p = doc.add_paragraph(note_text)
        if p.runs:                                             # add_paragraph rỗng → 0 run; guard IndexError
            p.runs[0].italic = True
    items = [it for it in _as_list(bundle.get("items")) if isinstance(it, dict)]
    doc.add_paragraph(f"Số câu/nhóm: {bundle.get('count', len(items))} · đạt QC: {bundle.get('count_qc_ok', '?')}")
    doc.add_paragraph("⚠️ Đề do AI sinh từ ngân hàng đối tác — GV tiếng Anh BẮT BUỘC soát/ký trước khi dùng.")

    answer_key = []
    for i, it in enumerate(items, 1):
        if renderer:
            renderer(doc, i, it, answer_key)
        else:                                     # skill lạ: dump gọn để không mất dữ liệu (fallback)
            _hdr(doc, it, i, "Item")
            dump = json.dumps({k: v for k, v in it.items() if k != "qc_issues"}, ensure_ascii=False)
            doc.add_paragraph(dump[:1500] + ("…" if len(dump) > 1500 else ""))
        doc.add_paragraph("")

    doc.add_page_break()
    doc.add_heading("ĐÁP ÁN (GV soát/ký)", level=1)
    for line in answer_key:
        doc.add_paragraph(line)
    if not answer_key:
        doc.add_paragraph("(Kỹ năng tự luận — chấm theo tiêu chí, không có đáp án cố định.)")
    doc.add_paragraph("\nGV duyệt (Đạt/Không): ________   Chữ ký: ________   Ngày: ________")

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


def render_bundle_file(bundle_path: str, out_path: str) -> str:
    """Đọc bundle JSON từ đĩa → render .docx (tiện cho CLI)."""
    with open(bundle_path, encoding="utf-8") as f:
        bundle = json.load(f)
    return render_bundle_docx(bundle, out_path)
