import os
import io
import docx

def create_b1_reading_docx(filepath):
    doc = docx.Document()
    
    # Header Table
    t0 = doc.add_table(rows=1, cols=2)
    t0.cell(0, 0).text = "BỘ GIÁO DỤC VÀ ĐÀO TẠO TRƯỜNG ĐẠI HỌC THÀNH ĐÔNG"
    t0.cell(0, 1).text = "ĐỀ THI TIẾNG ANH B1 CHÂU ÂU"
    
    # Set ID Table
    t1 = doc.add_table(rows=1, cols=1)
    t1.cell(0, 0).text = "Mã đề:   EB1.9999"
    
    doc.add_paragraph("PART ONE 		READING")
    
    # Section 1
    doc.add_paragraph("Section 1		Questions 1-10 (10 points)")
    doc.add_paragraph("Circle the letter next to the word or phrase which best completes each sentence (A, B, C or D)")
    for q in range(1, 11):
        doc.add_paragraph(f"{q}. Question content {q} _______")
        doc.add_paragraph(f"\tA. Option A\tB. Option B\t\tC. Option C\t\tD. Option D")
        
    # Section 2
    doc.add_paragraph("Section 2	Questions 11-15 (5 points)")
    doc.add_paragraph("Look at the text in each question. What does it say?")
    
    t2 = doc.add_table(rows=5, cols=3)
    for q in range(11, 16):
        r = q - 11
        t2.cell(r, 0).text = f"{q}."
        t2.cell(r, 1).text = f"Signboard text for question {q}"
        t2.cell(r, 2).text = f"A. Explanation A\nB. Explanation B\nC. Explanation C"
        
    # Section 3
    doc.add_paragraph("Section 3	Questions 16-20 (5 points)")
    doc.add_paragraph("Read the text and questions below.")
    doc.add_paragraph("Passage paragraph 1 of dentist text.")
    doc.add_paragraph("Passage paragraph 2. Sophia Ashley, Oxford.")
    
    for q in range(16, 21):
        doc.add_paragraph(f"{q}. Question dental {q}?")
        doc.add_paragraph("A. dentist A")
        doc.add_paragraph("B. dentist B")
        doc.add_paragraph("C. dentist C")
        doc.add_paragraph("D. dentist D")
        
    # Section 4
    doc.add_paragraph("Section 4	Questions 21-30 (10 points)")
    doc.add_paragraph("Read the text below and fill each of the blanks with ONE suitable word.")
    doc.add_paragraph("THE GORILLA")
    doc.add_paragraph("The gorilla is a shy creature. In (21) ……………, it is quite different. It stands up if it wants to (22) ………………… an enemy.")
    doc.add_paragraph("Gorillas are the largest and (23) …………….… powerful. Adult males (24) ……………… from 135 to 230 kg.")
    doc.add_paragraph("Females are smaller. (25) ……………………. males and females are strong. They (26) ……………… their days in search (27) …………….for food.")
    doc.add_paragraph("Unfortunately, few animals are (28) …………….. in the wild. This is because people cut forests in (29) …………………. gorillas live. If we want to save them, we (30) …………… take action.")
    
    doc.add_paragraph("PART TWO		WRITING")
    
    # Writing Section 1
    doc.add_paragraph("Section 1 (10 points)")
    doc.add_paragraph("Finish each of the following sentences...")
    doc.add_paragraph("Example:	I haven’t enjoyed myself so much")
    doc.add_paragraph("Answer:	It’s years since...")
    doc.add_paragraph("How is your surname spelt?")
    doc.add_paragraph("How do ……………………………………………………………………………………?")
    doc.add_paragraph("At the moment, they are cleaning Mr. Lazylion’s car.")
    doc.add_paragraph("At the moment Mr. Lazylion ……………………………..………………………………..")
    
    # Writing Section 2
    doc.add_paragraph("Section 2 (20 points)")
    doc.add_paragraph("You are Hoa Tran. Write a letter to your penfriend.")
    doc.add_paragraph("Dear friend, I’ve got flu. What should I do?")
    doc.add_paragraph("Now write a letter. You should write about 100 words.")
    doc.add_paragraph("--- The end ---")
    
    doc.save(filepath)


def create_b1_answer_key_docx(filepath):
    doc = docx.Document()
    
    doc.add_paragraph("MÃ PHÁCH: ………………; MÃ ĐỀ: EB1.9999")
    doc.add_paragraph("PART 1: READING")
    doc.add_paragraph("Câu 1 - 20:")
    
    # Table 1: Section 1, 2, 3 answers
    # 8 rows, 19 columns
    t1 = doc.add_table(rows=8, cols=19)
    # Row 0: Merged header blocks
    for c in range(9):
        t1.cell(0, c).text = "SECTION 1"
    t1.cell(0, 9).text = ""
    for c in range(10, 14):
        t1.cell(0, c).text = "SECTION 2"
    t1.cell(0, 14).text = ""
    for c in range(15, 19):
        t1.cell(0, c).text = "SECTION 3"
        
    # Row 1: Subheaders
    headers = ['Câu', 'Đáp án', 'Thang điểm', 'Điểm chấm', '']
    row1 = headers * 3 + ['Câu', 'Đáp án', 'Thang điểm', 'Điểm chấm']
    for c, val in enumerate(row1):
        t1.cell(1, c).text = val
        
    # Rows 2-6: Data
    sec1_ans = {1: "A", 2: "B", 3: "B", 4: "A", 5: "A", 6: "B", 7: "A", 8: "D", 9: "B", 10: "A"}
    sec2_ans = {11: "A", 12: "C", 13: "C", 14: "B", 15: "A"}
    sec3_ans = {16: "D", 17: "A", 18: "D", 19: "B", 20: "C"}
    
    for r in range(2, 7):
        offset = r - 2
        # Section 1 block 1 (Q1-5)
        t1.cell(r, 0).text = str(offset + 1)
        t1.cell(r, 1).text = sec1_ans[offset + 1]
        t1.cell(r, 2).text = "1"
        
        # Section 1 block 2 (Q6-10)
        t1.cell(r, 5).text = str(offset + 6)
        t1.cell(r, 6).text = sec1_ans[offset + 6]
        t1.cell(r, 7).text = "1"
        
        # Section 2 (Q11-15)
        t1.cell(r, 10).text = str(offset + 11)
        t1.cell(r, 11).text = sec2_ans[offset + 11]
        t1.cell(r, 12).text = "1"
        
        # Section 3 (Q16-20)
        t1.cell(r, 15).text = str(offset + 16)
        t1.cell(r, 16).text = sec3_ans[offset + 16]
        t1.cell(r, 17).text = "1"
        
    # Row 7: Totals
    t1.cell(7, 0).text = "Tổng"
    t1.cell(7, 2).text = "5"
    t1.cell(7, 5).text = "Tổng"
    t1.cell(7, 7).text = "5"
    t1.cell(7, 10).text = "Tổng"
    t1.cell(7, 12).text = "5"
    t1.cell(7, 15).text = "Tổng"
    t1.cell(7, 17).text = "5"
    
    doc.add_paragraph("Câu 21 - 30:")
    
    # Table 2: Section 4 answers
    # 8 rows, 9 columns
    t2 = doc.add_table(rows=8, cols=9)
    for c in range(9):
        t2.cell(0, c).text = "SECTION 4"
    headers_t2 = ['Câu', 'Đáp án', 'Thang điểm', 'Điểm chấm', '']
    row1_t2 = headers_t2 + ['Câu', 'Đáp án', 'Thang điểm', 'Điểm chấm']
    for c, val in enumerate(row1_t2):
        t2.cell(1, c).text = val
        
    sec4_ans = {
        21: "fact", 22: "frighten", 23: "most", 24: "weigh", 25: "Both",
        26: "spend", 27: "search", 28: "left", 29: "which", 30: "must"
    }
    for r in range(2, 7):
        offset = r - 2
        # Q21-25
        t2.cell(r, 0).text = str(offset + 21)
        t2.cell(r, 1).text = sec4_ans[offset + 21]
        t2.cell(r, 2).text = "1"
        
        # Q26-30
        t2.cell(r, 5).text = str(offset + 26)
        t2.cell(r, 6).text = sec4_ans[offset + 26]
        t2.cell(r, 7).text = "1"
        
    t2.cell(7, 0).text = "Tổng"
    t2.cell(7, 2).text = "5"
    t2.cell(7, 5).text = "Tổng"
    t2.cell(7, 7).text = "5"
    
    doc.add_paragraph("PART 2: WRITING")
    doc.add_paragraph("Tổng điểm bài thi: 60")
    
    doc.save(filepath)


def create_b1_listening_docx(filepath):
    import zlib
    import struct
    def _png_1px():
        def _chunk(tag, data):
            c = tag + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        sig = b"\x89PNG\r\n\x1a\n"
        ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        idat = zlib.compress(b"\x00\xff\x00\x00")
        return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")
    png_data = _png_1px()
    
    doc = docx.Document()
    
    # Add set_id in footer
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "Page 2 of              Mã đề thi LB1.2601"
    
    # Table 0: Q1-5
    t0 = doc.add_table(rows=10, cols=1)
    for i in range(5):
        t0.rows[2*i].cells[0].text = f"{i+1}. Which dish did Mark cook in the competition?"
        # Add image in even row cell
        cell = t0.rows[2*i+1].cells[0]
        cell.paragraphs[0].add_run().add_picture(io.BytesIO(png_data))
        
    # Table 1: Q6-11
    t1 = doc.add_table(rows=1, cols=1)
    t1.rows[0].cells[0].text = (
        "Name of Ben's organisation: (6) ………………\n"
        "Aim of course: Discovering (7) …………………..\n"
        "Closest course location for this group: (8) …………………….\n"
        "Length of course: (9) …………………. weeks\n"
        "• Cut up (10) ………………………\n"
        "• Make a (11) ………………………"
    )
    
    # Table 2: Q12-15
    t2 = doc.add_table(rows=1, cols=1)
    t2.rows[0].cells[0].text = (
        "Songs from musicals: (12) Susan ..............................\n"
        "Colour of clothes: (13) ..............................\n"
        "Map of building: available from the (14) ..............................\n"
        "Car park: costs £ (15) .............................. per day"
    )
    doc.save(filepath)


def create_b1_listening_key_docx(filepath):
    doc = docx.Document()
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = (
        "1. C\n2. B\n3. A\n4. B\n5. C\n"
        "6. Nature\n7. wildlife\n8. forest\n9. 12/twelve\n10. wood\n"
        "11. waysbury\n12. Brokley\n13. blue\n14. receptionist\n15. 3/ three"
    )
    doc.save(filepath)


def create_b1_speaking_docx(filepath):
    doc = docx.Document()
    lines = [
        "----------------------------------------",
        "Speaking card B1 2601",
        "Part 1: Introducing yourself (2 minutes)",
        "Part 2: Topic - Talk about your favourite season. (3 minutes)",
        "Part 3: Interaction between Teachers and Candidate (2 minutes)",
        "----------------------------------------",
        "Speaking card B1 2602",
        "Part 1: Introducing yourself (2 minutes)",
        "Part 2: Topic - Talk about your favourite room. (3 minutes)",
        "Part 3: Interaction between Teachers and Candidate (2 minutes)",
        "----------------------------------------"
    ]
    for line in lines:
        doc.add_paragraph(line)
    doc.save(filepath)


def main():
    # Make sure we are writing to the correct absolute directory
    dir_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "fixtures", "parser"))
    os.makedirs(dir_path, exist_ok=True)
    
    create_b1_reading_docx(os.path.join(dir_path, "B1_exam_sample.docx"))
    create_b1_answer_key_docx(os.path.join(dir_path, "B1_key_sample.docx"))
    create_b1_listening_docx(os.path.join(dir_path, "B1_listening_sample.docx"))
    create_b1_listening_key_docx(os.path.join(dir_path, "B1_listening_key_sample.docx"))
    create_b1_speaking_docx(os.path.join(dir_path, "B1_speaking_sample.docx"))

    print("Test fixtures generated successfully in:", dir_path)

if __name__ == "__main__":
    main()



